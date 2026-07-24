"""DOCX text extraction using python-docx."""

import hashlib
import zipfile
from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from pydantic import ValidationError

from app.models.document import ParsedDocument
from app.parser.base_parser import BaseParser, ParserException


class DocxParser(BaseParser):
    """Extract paragraph and table text from a Word document."""

    def parse(self, file_path: str | Path) -> ParsedDocument:
        """Return DOCX text, an estimated explicit page count, and source facts."""
        path = Path(file_path)
        try:
            content = path.read_bytes()
            document = Document(path)
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            table_rows = [
                " | ".join(cell.text for cell in row.cells)
                for table in document.tables
                for row in table.rows
            ]
            raw_text = "\n".join(item for item in [*paragraphs, *table_rows] if item)
            # DOCX does not contain rendered pagination; explicit page-break markers are reliable.
            explicit_breaks = document.element.xml.count('w:type="page"')
            metadata = {
                "format": "docx",
                "author": document.core_properties.author or None,
                "title": document.core_properties.title or None,
                "subject": document.core_properties.subject or None,
                "paragraph_count": len(paragraphs),
                "table_count": len(document.tables),
            }
            return ParsedDocument(
                filename=path.name,
                pages=max(1, explicit_breaks + 1),
                raw_text=raw_text,
                metadata={key: value for key, value in metadata.items() if value is not None},
                file_hash=hashlib.sha256(content).hexdigest(),
            )
        except (OSError, PackageNotFoundError, ValueError, KeyError, zipfile.BadZipFile, ValidationError) as error:
            raise ParserException(f"Unable to parse DOCX '{path.name}'.") from error
